from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from copy import deepcopy
import io
import openpyxl
from docx import Document
from pathlib import Path

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = BASE_DIR / "template.docx"
RECEIVED_MARK = "X"
SHOW_VALUES = {"received": "Received", "declined": "Declined", "not eligible": "Not Eligible"}
BLANK_VALUES = {"outstanding": ""}


def normalize(v):
    return "" if v is None else str(v).strip()


def mark_for_status(status, received_mark, not_eligible_text):
    s = normalize(status).lower()
    if s == "received":
        return received_mark
    if s == "not eligible":
        return not_eligible_text
    return ""


def iter_paragraphs_in_table(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for nested in cell.tables:
                yield from iter_paragraphs_in_table(nested)


def iter_all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        yield from iter_paragraphs_in_table(table)


def fill_template(group_name, payer_labels, statuses, template_bytes, received_mark, not_eligible_text):
    doc = Document(io.BytesIO(template_bytes))
    label_to_status = {}
    for i, label in enumerate(payer_labels):
        if label.strip() == "Payer":
            continue
        label_to_status[label] = mark_for_status(statuses[i], received_mark, not_eligible_text)

    for p in iter_all_paragraphs(doc):
        t = p.text.strip()

        if t.startswith("Group Name"):
            p.text = f"Group Name {group_name}"
            continue

        if t.startswith("Payer "):
            label = t.split("\t", 1)[0].strip()
            if label in label_to_status:
                p.text = f"{label}\t{label_to_status[label]}"

    return doc


@app.post("/generate")
async def generate(
    file: UploadFile = File(...),
    template: UploadFile | None = File(None),
    mark_symbol: str = Form(RECEIVED_MARK),
    not_eligible_text: str = Form("Not Eligible"),
):
    if template is not None:
        template_bytes = await template.read()
        if not template_bytes:
            raise HTTPException(status_code=400, detail="Uploaded template is empty")
    else:
        if not TEMPLATE_PATH.exists():
            raise HTTPException(
                status_code=500,
                detail=f"template.docx not found at {TEMPLATE_PATH}",
            )
        template_bytes = TEMPLATE_PATH.read_bytes()
    data = await file.read()
    ws = load_sheet(data)

    payer_letters = [normalize(ws.cell(2, c).value) for c in range(2, 10)]
    payer_labels = [f"Payer {x}" for x in payer_letters]

    out = None

    for r in range(4, ws.max_row + 1):
        group_name = normalize(ws.cell(r, 1).value)
        if not group_name:
            continue

        statuses = [ws.cell(r, c).value for c in range(2, 10)]
        temp = fill_template(
            group_name,
            payer_labels,
            statuses,
            template_bytes,
            mark_symbol or RECEIVED_MARK,
            not_eligible_text,
        )

        if out is None:
            out = temp
        else:
            out.add_page_break()
            for element in temp.element.body:
                if element.tag.endswith("sectPr"):
                    continue
                out.element.body.append(deepcopy(element))

    if out is None:
        out = Document()

    buf = io.BytesIO()
    out.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=Received_Plans_Report.docx"},
    )


@app.post("/generate-table")
async def generate_table(file: UploadFile = File(...)):
    data = await file.read()
    ws = load_sheet(data)
    payer_headers, rows = read_rows(ws)

    doc = Document()
    doc.add_heading("Group Status Report", level=1)

    table = doc.add_table(rows=len(rows) + 1, cols=len(payer_headers) + 1)
    table.style = "Table Grid"

    table.cell(0, 0).text = "Group"
    for j, h in enumerate(payer_headers, start=1):
        table.cell(0, j).text = h

    for i, (group, statuses) in enumerate(rows, start=1):
        table.cell(i, 0).text = group
        for j, v in enumerate(statuses, start=1):
            table.cell(i, j).text = display_status(v)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=Group_Status_Report.docx"},
    )
def display_status(v):
    s = normalize(v).lower()
    if s in SHOW_VALUES:
        return SHOW_VALUES[s]
    if s in BLANK_VALUES:
        return BLANK_VALUES[s]
    return "" if s == "" else normalize(v)


def load_sheet(data):
    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
    return wb.active


def read_rows(ws):
    payer_headers = []
    for c in range(2, 10):
        h = normalize(ws.cell(2, c).value)
        payer_headers.append(h if h else f"Col{c}")

    rows = []
    for r in range(4, ws.max_row + 1):
        group = normalize(ws.cell(r, 1).value)
        if not group:
            continue
        statuses = [ws.cell(r, c).value for c in range(2, 10)]
        rows.append((group, statuses))

    return payer_headers, rows


def payer_label_from_header(header, fallback_letter):
    h = normalize(header)
    return h if h else f"Payer {fallback_letter}"


@app.post("/generate-lines")
async def generate_lines(file: UploadFile = File(...)):
    data = await file.read()
    ws = load_sheet(data)

    letters = "ABCDEFGH"
    payer_labels = []
    for i, c in enumerate(range(2, 10)):
        header = ws.cell(2, c).value
        payer_labels.append(payer_label_from_header(header, letters[i]))

    doc = Document()
    first_group = True

    for r in range(4, ws.max_row + 1):
        group = normalize(ws.cell(r, 1).value)
        if not group:
            continue

        if not first_group:
            doc.add_page_break()
        first_group = False

        doc.add_paragraph(f"Group Name {group}")
        doc.add_paragraph("Received Plans")

        statuses = [ws.cell(r, c).value for c in range(2, 10)]
        for label, status in zip(payer_labels, statuses):
            value = normalize(status)
            line = f"{label} {value}".rstrip()
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=Group_Status_Report.docx"},
    )
