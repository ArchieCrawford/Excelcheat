from pathlib import Path
import openpyxl
from docx import Document

EXCEL_PATH = "Test.xlsx"
OUTPUT_DOCX = "Group_Status_Report.docx"

SHOW_VALUES = {"received": "Received", "declined": "Declined", "not eligible": "Not Eligible"}
BLANK_VALUES = {"outstanding": ""}

def norm(v):
    return "" if v is None else str(v).strip()

def display_status(v):
    s = norm(v).lower()
    if s in SHOW_VALUES:
        return SHOW_VALUES[s]
    if s in BLANK_VALUES:
        return BLANK_VALUES[s]
    return "" if s == "" else norm(v)

wb = openpyxl.load_workbook(EXCEL_PATH, data_only=True)
ws = wb.active

payer_headers = []
for c in range(2, 10):
    h = norm(ws.cell(2, c).value)
    payer_headers.append(h if h else f"Col{c}")

rows = []
for r in range(4, ws.max_row + 1):
    group = norm(ws.cell(r, 1).value)
    if not group:
        continue
    statuses = [ws.cell(r, c).value for c in range(2, 10)]
    rows.append((group, statuses))

doc = Document()
doc.add_heading("Group Status Report", level=1)

table = doc.add_table(rows=len(rows) + 1, cols=len(payer_headers) + 1)
table.style = "Table Grid"

table.cell(0, 0).text = "Group"
for j, h in enumerate(payer_headers, start=1):
    table.cell(0, j).text = h

for i, (group, statuses) in enumerate(rows, start=1):
    table.cell(i, 0).text = group
    for j, v in enumerate(statuses, start=1):
        table.cell(i, j).text = display_status(v)

doc.save(OUTPUT_DOCX)
print(f"Saved: {OUTPUT_DOCX}")